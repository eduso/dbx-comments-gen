# Databricks notebook source

# MAGIC %md
# MAGIC # Generador de Comentarios v4
# MAGIC ## Documentación Automática de Esquemas, Tablas y Columnas con IA
# MAGIC
# MAGIC Este notebook genera comentarios descriptivos orientados a **usuarios de negocio** para cada
# MAGIC **esquema**, **tabla** y **columna** de un catálogo de Unity Catalog, utilizando un modelo
# MAGIC fundacional de IA como motor de generación.
# MAGIC
# MAGIC ### Alcance
# MAGIC - Recorre todos los esquemas de un catálogo, o uno específico si se indica.
# MAGIC - Para cada esquema, genera un comentario de negocio y luego procesa todas sus tablas y columnas.
# MAGIC - Utiliza documentos proporcionados por el usuario (en `input/`) como contexto para el modelo.
# MAGIC - Opcionalmente, toma una muestra de datos reales de cada tabla para enriquecer el contexto.
# MAGIC
# MAGIC ### Parámetros del Notebook
# MAGIC | Parámetro | Descripción | Default |
# MAGIC |-----------|-------------|---------|
# MAGIC | `catalog_name` | Catálogo de Unity Catalog a procesar | `main_eduardo_sojo` |
# MAGIC | `schema_name` | Esquema específico. Si está vacío, procesa todos los esquemas del catálogo | _(vacío)_ |
# MAGIC | `model_endpoint` | Nombre del endpoint del modelo fundacional en Foundation Model API | `databricks-claude-sonnet-4-5` |
# MAGIC | `results_catalog` | Catálogo donde se persisten los resultados de las ejecuciones | `main_eduardo_sojo` |
# MAGIC | `results_schema` | Esquema donde se persisten los resultados de las ejecuciones | `ai_comments_generation` |
# MAGIC | `enable_sampling` | Habilitar muestreo de datos reales como contexto adicional (`yes`/`no`) | `no` |
# MAGIC | `sampling_pct` | Porcentaje de muestreo cuando la tabla tiene más de 500 registros (1-100) | `10` |
# MAGIC
# MAGIC ### Insumos
# MAGIC El notebook lee el archivo `input/mapping.md` para descubrir los documentos de contexto y su propósito.
# MAGIC Cada archivo listado en el mapping se carga y se prioriza dinámicamente según su relevancia
# MAGIC para el esquema/tabla que se esté procesando en cada momento.
# MAGIC
# MAGIC ### Tablas de Resultados
# MAGIC - `{results_catalog}.{results_schema}.ejecuciones` — Registro de cada ejecución del proceso
# MAGIC - `{results_catalog}.{results_schema}.resultados` — Comentarios generados por item:
# MAGIC   - Esquemas → `nombre_tabla = '__esquema__'`, `nombre_columna = '__esquema__'`
# MAGIC   - Tablas   → `nombre_columna = '__tabla__'`
# MAGIC   - Columnas → `nombre_columna = <nombre_técnico>`

# COMMAND ----------

# MAGIC %md
# MAGIC ## 1. Instalación de Dependencias
# MAGIC
# MAGIC Se instalan las librerías necesarias:
# MAGIC - `python-docx`: Para leer archivos .docx proporcionados como insumo
# MAGIC - `mlflow`: Para interactuar con Foundation Model API a través del deploy client

# COMMAND ----------

# MAGIC %pip install python-docx mlflow -q
# MAGIC dbutils.library.restartPython()

# COMMAND ----------

# MAGIC %md
# MAGIC ## 2. Lectura de Parámetros
# MAGIC
# MAGIC Se definen los widgets del notebook y se leen los valores proporcionados por el usuario.
# MAGIC Estos parámetros controlan:
# MAGIC - **Qué procesar**: catálogo y esquema objetivo
# MAGIC - **Con qué modelo**: endpoint del modelo fundacional
# MAGIC - **Dónde guardar**: catálogo y esquema de resultados
# MAGIC - **Sampling**: si se habilita y con qué porcentaje

# COMMAND ----------

# -- Definición de widgets --
dbutils.widgets.text("catalog_name",    "main_eduardo_sojo",            "Catálogo a procesar")
dbutils.widgets.text("schema_name",     "",                             "Esquema (vacío = todo el catálogo)")
dbutils.widgets.text("model_endpoint",  "databricks-claude-sonnet-4-5", "Modelo fundacional")
dbutils.widgets.text("results_catalog", "main_eduardo_sojo",            "Catálogo de resultados")
dbutils.widgets.text("results_schema",  "ai_comments_generation",       "Esquema de resultados")
dbutils.widgets.dropdown("enable_sampling", "no", ["no", "yes"],        "Sampling de datos")
dbutils.widgets.text("sampling_pct",    "10",                           "Porcentaje de sampling (1-100)")

# -- Lectura de parámetros --
CATALOG_NAME    = dbutils.widgets.get("catalog_name").strip()
SCHEMA_NAME     = dbutils.widgets.get("schema_name").strip()
MODEL_ENDPOINT  = dbutils.widgets.get("model_endpoint").strip()
RESULTS_CATALOG = dbutils.widgets.get("results_catalog").strip()
RESULTS_SCHEMA  = dbutils.widgets.get("results_schema").strip()
ENABLE_SAMPLING = dbutils.widgets.get("enable_sampling").strip().lower() == "yes"

# -- Validar porcentaje de sampling --
try:
    SAMPLING_PCT = int(dbutils.widgets.get("sampling_pct").strip())
    if not 1 <= SAMPLING_PCT <= 100:
        raise ValueError
except ValueError:
    raise ValueError("El parámetro 'sampling_pct' debe ser un entero entre 1 y 100.")

# -- Resumen de parámetros --
print("=" * 60)
print("PARÁMETROS DE EJECUCIÓN")
print("=" * 60)
print(f"  Catálogo a procesar : {CATALOG_NAME}")
print(f"  Esquema             : {SCHEMA_NAME or '(todos los esquemas del catálogo)'}")
print(f"  Modelo              : {MODEL_ENDPOINT}")
print(f"  Resultados en       : {RESULTS_CATALOG}.{RESULTS_SCHEMA}")
print(f"  Sampling habilitado : {ENABLE_SAMPLING}")
print(f"  Porcentaje sampling : {SAMPLING_PCT}%")
print("=" * 60)

# COMMAND ----------

# MAGIC %md
# MAGIC ## 3. Configuración del Entorno
# MAGIC
# MAGIC Se inicializa:
# MAGIC - El **logger** para registro detallado de cada etapa del proceso
# MAGIC - Las **rutas de tablas de control** donde se persisten ejecuciones y resultados
# MAGIC - La **ruta del directorio de insumos** (`input/`) relativa a la ubicación del notebook
# MAGIC - El **deploy client** de MLflow para invocar el modelo fundacional

# COMMAND ----------

import uuid
import os
import logging
from datetime import datetime, timezone
from mlflow.deployments import get_deploy_client

# ── Configuración del logger ────────────────────────────────────────────────
# Se usa un logger con nombre propio para poder filtrar los mensajes en el
# output del notebook y diferenciarlos de otros logs del sistema.
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("comments_generator_v4")

# ── Nombres completamente calificados de las tablas de control ──────────────
EXEC_TABLE    = f"{RESULTS_CATALOG}.{RESULTS_SCHEMA}.ejecuciones"
RESULTS_TABLE = f"{RESULTS_CATALOG}.{RESULTS_SCHEMA}.resultados"

# ── Ruta del directorio de insumos ──────────────────────────────────────────
# El notebook está en src/notebooks/, por lo que subimos dos niveles para
# llegar a la raíz del proyecto y luego acceder a input/
notebook_path = (
    dbutils.notebook.entry_point
    .getDbutils().notebook().getContext()
    .notebookPath().get()
)
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(notebook_path)))
INPUT_DIR    = f"/Workspace{PROJECT_ROOT}/input"

# ── Cliente de MLflow para Foundation Model API ─────────────────────────────
deploy_client = get_deploy_client("databricks")

logger.info(f"Endpoint         : {MODEL_ENDPOINT}")
logger.info(f"Tabla ejecuciones: {EXEC_TABLE}")
logger.info(f"Tabla resultados : {RESULTS_TABLE}")
logger.info(f"Directorio input : {INPUT_DIR}")

# COMMAND ----------

# MAGIC %md
# MAGIC ## 4. Preparación del Esquema de Resultados
# MAGIC
# MAGIC Crea automáticamente el esquema y las tablas de control si no existen.
# MAGIC Esto permite al usuario desplegar la solución apuntando a cualquier
# MAGIC catálogo/esquema sin necesidad de ejecutar scripts DDL por separado.
# MAGIC
# MAGIC **Tablas creadas:**
# MAGIC - `ejecuciones`: Registra cada ejecución con su ID, fecha, estado y resultado
# MAGIC - `resultados`: Almacena cada comentario generado vinculado a su ejecución

# COMMAND ----------

logger.info("ETAPA: Preparación del esquema de resultados")
logger.info(f"  Verificando existencia de {RESULTS_CATALOG}.{RESULTS_SCHEMA}...")

# -- Crear esquema si no existe --
spark.sql(f"""
    CREATE SCHEMA IF NOT EXISTS {RESULTS_CATALOG}.{RESULTS_SCHEMA}
    COMMENT 'Esquema para la generación automática de comentarios usando IA'
""")
logger.info(f"  ✓ Esquema {RESULTS_CATALOG}.{RESULTS_SCHEMA} verificado")

# -- Crear tabla de ejecuciones si no existe --
spark.sql(f"""
    CREATE TABLE IF NOT EXISTS {EXEC_TABLE} (
        id_ejecucion    STRING        NOT NULL COMMENT 'GUID único de la ejecución',
        fecha_ejecucion TIMESTAMP     COMMENT 'Fecha y hora de la ejecución (UTC)',
        estado          VARCHAR(50)   COMMENT 'Estado: INICIADO, EN_PROCESO, COMPLETADO, COMPLETADO_CON_ERRORES, ERROR',
        resultado       VARCHAR(4000) COMMENT 'Detalle del resultado final de la ejecución',
        CONSTRAINT pk_ejecuciones PRIMARY KEY (id_ejecucion)
    )
    COMMENT 'Registro de cada ejecución del proceso de generación de comentarios'
    TBLPROPERTIES ('delta.enableChangeDataFeed' = 'true')
""")
logger.info(f"  ✓ Tabla {EXEC_TABLE} verificada")

# -- Crear tabla de resultados si no existe --
spark.sql(f"""
    CREATE TABLE IF NOT EXISTS {RESULTS_TABLE} (
        id_resultado    BIGINT        GENERATED ALWAYS AS IDENTITY COMMENT 'ID autoincrementable del resultado',
        id_ejecucion    STRING        NOT NULL COMMENT 'Referencia a la ejecución que generó este resultado',
        fecha_resultado TIMESTAMP     COMMENT 'Fecha en que se generó el comentario (UTC)',
        nombre_esquema  VARCHAR(255)  COMMENT 'Esquema de la tabla procesada',
        nombre_tabla    VARCHAR(255)  COMMENT 'Tabla procesada (o __esquema__ para comentarios de esquema)',
        nombre_columna  VARCHAR(255)  COMMENT 'Columna procesada (o __tabla__/__esquema__ para comentarios de tabla/esquema)',
        comentario      VARCHAR(4000) COMMENT 'Comentario generado por IA',
        CONSTRAINT pk_resultados PRIMARY KEY (id_resultado),
        CONSTRAINT fk_ejecucion  FOREIGN KEY (id_ejecucion)
            REFERENCES {EXEC_TABLE} (id_ejecucion)
    )
    COMMENT 'Comentarios generados por IA por columna de cada tabla procesada'
    TBLPROPERTIES ('delta.enableChangeDataFeed' = 'true')
""")
logger.info(f"  ✓ Tabla {RESULTS_TABLE} verificada")
logger.info("  Esquema de resultados listo")

# COMMAND ----------

# MAGIC %md
# MAGIC ## 5. Funciones de Persistencia
# MAGIC
# MAGIC Funciones auxiliares para interactuar con las tablas de control:
# MAGIC - `_now()`: Timestamp UTC actual para registros SQL
# MAGIC - `_esc()`: Escape de comillas simples para prevenir SQL injection
# MAGIC - `insert_ejecucion()` / `update_ejecucion()`: Gestión del ciclo de vida de una ejecución
# MAGIC - `insert_resultado()`: Inserción de un comentario generado

# COMMAND ----------

def _now() -> str:
    """
    Retorna la fecha/hora UTC actual en formato ISO 8601.
    Se usa como valor para campos TIMESTAMP en las inserciones SQL.
    """
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")


def _esc(value: str) -> str:
    """
    Escapa comillas simples en un string para uso seguro en sentencias SQL.
    Previene errores de sintaxis y posibles inyecciones SQL.
    """
    return (value or "").replace("'", "''")


# ── Gestión de ejecuciones ──────────────────────────────────────────────────

def insert_ejecucion(exec_id: str, estado: str) -> None:
    """
    Registra una nueva ejecución en estado inicial.
    Se invoca una sola vez al inicio del proceso con estado 'INICIADO'.
    """
    spark.sql(f"""
        INSERT INTO {EXEC_TABLE}
            (id_ejecucion, fecha_ejecucion, estado, resultado)
        VALUES
            ('{exec_id}', TIMESTAMP '{_now()}', '{_esc(estado)}', NULL)
    """)
    logger.info(f"  Ejecución registrada: {exec_id} [{estado}]")


def update_ejecucion(exec_id: str, estado: str, resultado: str = None) -> None:
    """
    Actualiza el estado y opcionalmente el resultado de una ejecución existente.
    Se invoca múltiples veces durante el proceso para reflejar el progreso.
    """
    resultado_sql = f"'{_esc(resultado)}'" if resultado is not None else "NULL"
    spark.sql(f"""
        UPDATE {EXEC_TABLE}
        SET estado = '{_esc(estado)}',
            fecha_ejecucion = TIMESTAMP '{_now()}',
            resultado = {resultado_sql}
        WHERE id_ejecucion = '{exec_id}'
    """)


# ── Persistencia de resultados ──────────────────────────────────────────────

def insert_resultado(
    exec_id: str, nombre_esquema: str, nombre_tabla: str,
    nombre_columna: str, comentario: str,
) -> None:
    """
    Inserta un comentario generado en la tabla de resultados.

    Convenciones:
    - Comentarios de esquema: nombre_tabla='__esquema__', nombre_columna='__esquema__'
    - Comentarios de tabla:   nombre_columna='__tabla__'
    - Comentarios de columna: nombre_columna=<nombre técnico de la columna>
    """
    spark.sql(f"""
        INSERT INTO {RESULTS_TABLE}
            (id_ejecucion, fecha_resultado, nombre_esquema, nombre_tabla, nombre_columna, comentario)
        VALUES (
            '{exec_id}', TIMESTAMP '{_now()}',
            '{_esc(nombre_esquema)}', '{_esc(nombre_tabla)}',
            '{_esc(nombre_columna)}', '{_esc(comentario)}'
        )
    """)

# COMMAND ----------

# MAGIC %md
# MAGIC ## 6. Carga de Insumos desde input/mapping.md
# MAGIC
# MAGIC El proceso de carga sigue estos pasos:
# MAGIC 1. Lee `input/mapping.md` y parsea las entradas en formato `nombre_archivo.ext: descripción`
# MAGIC 2. Para cada archivo referenciado, lo carga según su extensión (.docx, .tsv, .csv, .txt, .xlsx, etc.)
# MAGIC 3. Construye un diccionario indexado por nombre de archivo con su contenido y metadata
# MAGIC
# MAGIC Si el archivo `mapping.md` no existe, el proceso continúa sin contexto de insumos
# MAGIC (los prompts se generan solo con la metadata del catálogo).

# COMMAND ----------

import pandas as pd

logger.info("ETAPA: Carga de insumos desde mapping.md")


def _clean(value) -> str:
    """Limpia valores de pandas: convierte NaN/None a cadena vacía."""
    s = str(value).strip() if value is not None else ""
    return "" if s.lower() in ("nan", "none", "") else s


def parse_mapping_file(input_dir: str) -> list:
    """
    Lee input/mapping.md y extrae las entradas de archivos con su descripción.

    Formato esperado por línea:
        `nombre_archivo.ext`: Descripción de uso
        nombre_archivo.ext: Descripción de uso

    Ignora líneas vacías y encabezados markdown (#).

    Returns:
        Lista de dicts con keys 'filename' y 'description'.
    """
    mapping_path = f"{input_dir}/mapping.md"
    entries = []
    try:
        with open(mapping_path, "r", encoding="utf-8") as f:
            content = f.read()
        logger.info(f"  Cargado mapping.md ({len(content)} chars)")

        for line in content.split("\n"):
            line = line.strip()
            # Ignorar líneas vacías y encabezados markdown
            if not line or line.startswith("#"):
                continue
            # Parsear formato "nombre: descripción"
            if ":" in line:
                parts = line.split(":", 1)
                filename = parts[0].strip().strip("`").strip()
                description = parts[1].strip()
                # Validar que sea un nombre de archivo (tiene extensión)
                if filename and description and "." in filename:
                    entries.append({
                        "filename": filename,
                        "description": description,
                    })

        logger.info(f"  Archivos mapeados: {len(entries)}")
        for e in entries:
            logger.info(f"    - {e['filename']}: {e['description'][:80]}")

    except FileNotFoundError:
        logger.warning(f"  No se encontró {mapping_path} — se ejecutará sin insumos de contexto")
    except Exception as e:
        logger.error(f"  Error leyendo mapping.md: {e}")

    return entries


def load_file_content(filepath: str) -> str:
    """
    Carga el contenido de un archivo como texto según su extensión.

    Extensiones soportadas:
    - .docx: Extrae texto de todos los párrafos (requiere python-docx)
    - .tsv/.csv: Lee como DataFrame y convierte a texto tabular (máx 200 filas)
    - .txt/.md/.json/.yaml: Lectura directa como texto plano
    - .xls/.xlsx: Lee como DataFrame y convierte a texto tabular (máx 200 filas)
    - Otros: Intenta lectura como texto plano
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".docx":
        from docx import Document
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    elif ext in (".tsv", ".csv"):
        sep = "\t" if ext == ".tsv" else ","
        df = pd.read_csv(filepath, sep=sep, encoding="utf-8")
        return df.to_string(index=False, max_rows=200)

    elif ext in (".txt", ".md", ".json", ".yaml", ".yml"):
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()

    elif ext in (".xls", ".xlsx"):
        df = pd.read_excel(filepath)
        return df.to_string(index=False, max_rows=200)

    else:
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        except Exception:
            logger.warning(f"  No se pudo leer {filepath} — formato no soportado")
            return ""


def load_all_insumos(input_dir: str) -> dict:
    """
    Carga todos los insumos referenciados en mapping.md.

    Returns:
        Dict indexado por nombre de archivo. Cada valor contiene:
        - description: Propósito del archivo según mapping.md
        - content: Contenido del archivo como texto
        - size: Tamaño del contenido en caracteres
    """
    mapping = parse_mapping_file(input_dir)
    insumos = {}

    for entry in mapping:
        filepath = f"{input_dir}/{entry['filename']}"
        try:
            content = load_file_content(filepath)
            insumos[entry["filename"]] = {
                "description": entry["description"],
                "content": content,
                "size": len(content),
            }
            logger.info(f"    ✓ Cargado: {entry['filename']} ({len(content)} chars)")
        except FileNotFoundError:
            logger.warning(f"    ⚠ Archivo no encontrado: {entry['filename']}")
        except Exception as e:
            logger.error(f"    ✗ Error cargando {entry['filename']}: {e}")

    return insumos


# -- Ejecutar carga de insumos --
INSUMOS = load_all_insumos(INPUT_DIR)
logger.info(f"  Total insumos cargados: {len(INSUMOS)}")
for name, data in INSUMOS.items():
    logger.info(f"    {name}: {data['size']} chars — {data['description'][:60]}")

# COMMAND ----------

# MAGIC %md
# MAGIC ## 7. Motor de Contexto Dinámico
# MAGIC
# MAGIC El contexto para el modelo se construye dinámicamente para cada esquema/tabla:
# MAGIC 1. Se calcula un **score de relevancia** para cada insumo basado en:
# MAGIC    - Coincidencia del nombre del esquema o tabla en el contenido/nombre del archivo
# MAGIC    - Tipo de archivo (documentación .docx pesa más que datos .tsv)
# MAGIC    - Palabras clave en la descripción del mapping
# MAGIC 2. Se ordenan por relevancia descendente
# MAGIC 3. Se incluyen en el prompt respetando un **límite de ~30K caracteres** (~7,500 tokens)
# MAGIC 4. Si un archivo no cabe completo, se trunca. Si no cabe nada, se omite.
# MAGIC
# MAGIC Esto garantiza que el modelo siempre reciba el contexto más relevante sin exceder su ventana.

# COMMAND ----------

# Límite de caracteres para el contexto de insumos en cada prompt.
# ~30K chars ≈ 7,500 tokens, dejando espacio para el resto del prompt.
MAX_CONTEXT_CHARS = 30_000


def _score_relevance(filename: str, description: str, content: str,
                     schema_name: str, table_name: str = "") -> float:
    """
    Calcula un score de relevancia de un insumo para el esquema/tabla dado.

    Criterios de puntuación:
    - Base: 1.0 (todo insumo tiene algo de relevancia)
    - +10.0 si el nombre del esquema aparece en el contenido/nombre/descripción
    - +15.0 si el nombre de la tabla aparece en el contenido/nombre/descripción
    - +3.0 para .docx (definiciones de negocio)
    - +2.0 para .tsv (datos tabulares / ejemplos)
    - +1.5 para .md/.txt
    - +1.0 por cada palabra clave relevante encontrada en la descripción

    Returns:
        Float con el score. Mayor = más relevante.
    """
    score = 1.0
    text_lower = (description + " " + filename + " " + content[:500]).lower()
    schema_lower = schema_name.lower()
    table_lower = table_name.lower() if table_name else ""

    # Coincidencia con esquema y tabla
    if schema_lower and schema_lower in text_lower:
        score += 10.0
    if table_lower and table_lower in text_lower:
        score += 15.0

    # Bonus por tipo de archivo
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".docx":
        score += 3.0
    elif ext == ".tsv":
        score += 2.0
    elif ext in (".md", ".txt"):
        score += 1.5

    # Bonus por palabras clave en la descripción
    keywords = ["definic", "ejemplo", "comentario", "tabla", "columna",
                "campo", "negocio", "instruccion", "regla"]
    for kw in keywords:
        if kw in description.lower():
            score += 1.0

    return score


def build_dynamic_context(
    insumos: dict,
    schema_name: str,
    table_name: str = "",
    max_chars: int = MAX_CONTEXT_CHARS,
) -> str:
    """
    Construye el bloque de contexto para el prompt priorizando insumos relevantes.

    Proceso:
    1. Calcula score de relevancia para cada insumo
    2. Ordena por score descendente
    3. Incluye insumos hasta alcanzar el límite de caracteres
    4. Trunca el último archivo si no cabe completo (mínimo 500 chars útiles)

    Args:
        insumos: Dict de insumos cargados desde mapping.md
        schema_name: Nombre del esquema actual (para calcular relevancia)
        table_name: Nombre de la tabla actual (opcional, para mayor precisión)
        max_chars: Límite máximo de caracteres para el contexto

    Returns:
        String con el contexto concatenado, listo para insertar en el prompt.
    """
    if not insumos:
        return ""

    # Calcular relevancia de cada insumo
    scored = []
    for filename, data in insumos.items():
        score = _score_relevance(
            filename, data["description"], data["content"],
            schema_name, table_name,
        )
        scored.append((score, filename, data))

    # Ordenar por relevancia descendente
    scored.sort(key=lambda x: x[0], reverse=True)

    # Construir contexto respetando el límite de caracteres
    context_parts = []
    chars_used = 0

    for score, filename, data in scored:
        header = f"\n--- Insumo: {filename} (Propósito: {data['description']}) ---\n"
        content = data["content"]
        entry_size = len(header) + len(content)

        if chars_used + entry_size > max_chars:
            # Intentar incluir un fragmento truncado del archivo
            remaining = max_chars - chars_used - len(header) - 50
            if remaining > 500:
                content = content[:remaining] + "\n[... contenido truncado por límite de contexto]"
                context_parts.append(header + content)
                chars_used += len(header) + len(content)
                logger.info(f"    Contexto truncado: {filename} ({remaining} de {data['size']} chars)")
            else:
                logger.info(f"    Contexto omitido por límite: {filename} (score={score:.1f})")
            break
        else:
            context_parts.append(header + content)
            chars_used += entry_size

    logger.info(f"    Contexto construido: {chars_used} chars, {len(context_parts)} insumo(s) incluidos")
    return "\n".join(context_parts)

# COMMAND ----------

# MAGIC %md
# MAGIC ## 8. Sampling de Datos
# MAGIC
# MAGIC Cuando `enable_sampling` está activado, se toma una muestra aleatoria de cada tabla
# MAGIC para incluirla como contexto adicional en el prompt. Esto permite al modelo entender
# MAGIC mejor el contenido real de la tabla y generar comentarios más precisos.
# MAGIC
# MAGIC **Reglas de muestreo:**
# MAGIC - Si la tabla tiene **≤ 500 registros**: se toman todos
# MAGIC - Si tiene **> 500 registros**: se toma el porcentaje definido en `sampling_pct`
# MAGIC - La muestra se limita a **30 filas x 20 columnas** en el prompt para no exceder el contexto
# MAGIC - El muestreo es **aleatorio** (`ORDER BY RAND()`)

# COMMAND ----------

def get_table_sample(catalog: str, schema: str, table: str,
                     sampling_pct: int = SAMPLING_PCT) -> str:
    """
    Obtiene una muestra aleatoria de la tabla para usar como contexto en el prompt.

    Args:
        catalog: Nombre del catálogo
        schema: Nombre del esquema
        table: Nombre de la tabla
        sampling_pct: Porcentaje de muestreo para tablas con más de 500 registros (1-100)

    Returns:
        String con la muestra en formato tabular legible, o un mensaje descriptivo
        si la tabla está vacía o hay un error.
    """
    fqn = f"{catalog}.{schema}.{table}"
    try:
        # Obtener el conteo total de registros
        count_row = spark.sql(f"SELECT COUNT(*) AS cnt FROM {fqn}").collect()
        total_rows = count_row[0]["cnt"]

        if total_rows == 0:
            logger.info(f"      Sampling {table}: tabla vacía")
            return "(tabla vacía — sin registros)"

        # Determinar tamaño de la muestra según las reglas
        if total_rows <= 500:
            # Tablas pequeñas: tomar todos los registros
            sample_size = total_rows
            logger.info(f"      Sampling {table}: {total_rows} registros (todos — tabla pequeña)")
        else:
            # Tablas grandes: aplicar el porcentaje parametrizado
            fraction = sampling_pct / 100.0
            sample_size = max(int(total_rows * fraction), 1)
            logger.info(f"      Sampling {table}: {sample_size} de {total_rows} registros ({sampling_pct}%)")

        # Ejecutar el muestreo aleatorio
        sample_df = spark.sql(f"SELECT * FROM {fqn} ORDER BY RAND() LIMIT {sample_size}")

        # Convertir a pandas y limitar dimensiones para el prompt
        pdf = sample_df.toPandas()
        pdf = pdf.iloc[:30, :20]  # Máx 30 filas, 20 columnas
        return pdf.to_string(index=False, max_colwidth=50)

    except Exception as e:
        logger.warning(f"      ⚠ Error al obtener muestra de {fqn}: {e}")
        return f"(error al obtener muestra: {str(e)[:100]})"

# COMMAND ----------

# MAGIC %md
# MAGIC ## 9. Funciones de Generación de Comentarios
# MAGIC
# MAGIC Cada función construye un prompt especializado para el tipo de item a documentar:
# MAGIC - `generate_schema_comment()`: Comentario a nivel de esquema (dominio de negocio)
# MAGIC - `generate_table_comment()`: Comentario a nivel de tabla (granularidad, usos, reglas)
# MAGIC - `generate_column_comment()`: Comentario a nivel de columna (propósito, valores, lógica)
# MAGIC
# MAGIC Todos los prompts incluyen:
# MAGIC - El contexto dinámico construido a partir de los insumos
# MAGIC - Metadata del catálogo (comentarios existentes del esquema/tabla)
# MAGIC - Opcionalmente, la muestra de datos reales si el sampling está activo

# COMMAND ----------

def _call_model(prompt: str, max_tokens: int = 500) -> str:
    """
    Invoca el modelo fundacional a través de Foundation Model API.

    Args:
        prompt: Texto completo del prompt a enviar al modelo
        max_tokens: Máximo de tokens en la respuesta (default: 500)

    Returns:
        Texto de la respuesta del modelo, limpio de espacios.
    """
    response = deploy_client.predict(
        endpoint=MODEL_ENDPOINT,
        inputs={
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": max_tokens,
            "temperature": 0.2,
        },
    )
    return response["choices"][0]["message"]["content"].strip()


def generate_schema_comment(
    schema_name: str,
    tables_in_schema: list,
    context: str,
) -> str:
    """
    Genera un comentario de negocio para un esquema.

    El prompt incluye el nombre del esquema y la lista de tablas que contiene
    para que el modelo pueda inferir el dominio de negocio que cubre.
    """
    tables_list = ", ".join(tables_in_schema[:30])
    prompt = (
        "Eres un experto en documentación de datos orientada a usuarios de negocio "
        "de una entidad financiera.\n"
        + context
        + f"\nGenera una definición de negocio para el siguiente esquema:"
        + f"\n\n- Esquema: {schema_name}"
        + f"\n- Tablas contenidas: {tables_list}\n\n"
        + "La definición debe:\n"
        + "- Explicar el propósito del esquema y qué dominio de negocio cubre.\n"
        + "- Mencionar los principales tipos de datos que contiene.\n"
        + "- Estar escrita en español, lenguaje claro para usuarios de negocio.\n"
        + "- Tener como máximo 500 caracteres.\n"
        + "- Responde ÚNICAMENTE con la definición, sin texto adicional, sin comillas."
    )
    return _call_model(prompt)


def generate_table_comment(
    schema_name: str,
    schema_comment: str,
    table_name: str,
    context: str,
    sample_data: str = "",
) -> str:
    """
    Genera un comentario de negocio para una tabla.

    El prompt incluye:
    - Contexto del esquema padre (si tiene comentario)
    - Insumos dinámicos relevantes para esta tabla
    - Muestra de datos reales (si sampling está activo)
    """
    # Bloque de contexto del esquema
    context_block = ""
    if schema_comment:
        context_block = f"\nContexto del esquema '{schema_name}': {schema_comment}\n"

    # Bloque de muestra de datos (si aplica)
    sample_block = ""
    if sample_data:
        sample_block = (
            f"\nMuestra de datos reales de la tabla (para entender su contenido):\n"
            f"```\n{sample_data[:3000]}\n```\n"
        )

    prompt = (
        "Eres un experto en documentación de datos orientada a usuarios de negocio "
        "de una entidad financiera.\n"
        + context
        + context_block
        + sample_block
        + f"\nGenera una definición de negocio clara y completa para la siguiente tabla:"
        + f"\n\n- Tabla: {table_name}\n\n"
        + "La definición debe:\n"
        + "- Explicar a qué nivel se encuentra la información (ej: cliente, transacción, producto).\n"
        + "- Indicar los principales usos de la información en el negocio.\n"
        + "- Mencionar reglas o lógicas de negocio relevantes si el nombre lo sugiere.\n"
        + "- Estar escrita en español, en lenguaje claro para usuarios de negocio.\n"
        + "- Tener como máximo 500 caracteres.\n"
        + "- Responde ÚNICAMENTE con la definición, sin texto adicional, sin comillas."
    )
    return _call_model(prompt)


def generate_column_comment(
    schema_name: str,
    schema_comment: str,
    table_name: str,
    table_comment: str,
    column_name: str,
    data_type: str,
    context: str,
    sample_data: str = "",
) -> str:
    """
    Genera un comentario de negocio para una columna.

    El prompt incluye contexto jerárquico (esquema → tabla → columna) para que
    el modelo entienda el dominio y genere un comentario coherente con el resto.
    """
    # Bloque de contexto jerárquico (esquema + tabla)
    context_block = ""
    if schema_comment:
        context_block += f"\nContexto del esquema: {schema_comment}"
    if table_comment:
        context_block += f"\nContexto de la tabla '{table_name}': {table_comment}"
    context_block += "\n" if context_block else ""

    # Bloque de muestra de datos (si aplica)
    sample_block = ""
    if sample_data:
        sample_block = (
            f"\nMuestra de datos reales de la tabla (para entender el contenido de la columna):\n"
            f"```\n{sample_data[:3000]}\n```\n"
        )

    prompt = (
        "Eres un experto en documentación de datos orientada a usuarios de negocio "
        "de una entidad financiera.\n"
        + context
        + context_block
        + sample_block
        + f"\nGenera una definición de negocio clara y completa para la siguiente columna:"
        + f"\n\n- Tabla  : {table_name}"
        + f"\n- Columna: {column_name}"
        + f"\n- Tipo   : {data_type}\n\n"
        + "La definición debe:\n"
        + "- Describir el propósito o el contenido del campo para un usuario de negocio.\n"
        + "- Usar los nombres funcionales y sinónimos si los insumos los proporcionan.\n"
        + "- Incluir reglas o lógicas de negocio del campo si el nombre o la fórmula lo sugiere.\n"
        + "- Explicar los posibles valores si es catálogo, indicador (flg/ind/tip) o código.\n"
        + "- Estar escrita en español, en lenguaje claro y accesible.\n"
        + "- Tener como máximo 300 caracteres.\n"
        + "- Responde ÚNICAMENTE con la definición, sin texto adicional, sin comillas."
    )
    return _call_model(prompt, max_tokens=400)

# COMMAND ----------

# MAGIC %md
# MAGIC ## 10. Descubrimiento de Esquemas y Tablas
# MAGIC
# MAGIC Consulta `information_schema` para determinar qué esquemas y tablas procesar:
# MAGIC - Si `schema_name` está definido, procesa solo ese esquema
# MAGIC - Si está vacío, descubre todos los esquemas del catálogo (excluyendo `information_schema` y `default`)
# MAGIC - Para cada esquema, descubre todas sus tablas y columnas con sus tipos de dato

# COMMAND ----------

logger.info("ETAPA: Descubrimiento de esquemas y tablas")


def discover_schemas(catalog: str, schema_filter: str = "") -> list:
    """
    Descubre los esquemas a procesar en el catálogo.

    Args:
        catalog: Nombre del catálogo en Unity Catalog
        schema_filter: Si se especifica, retorna solo ese esquema. Si está vacío,
                      retorna todos los esquemas excluyendo los del sistema.

    Returns:
        Lista de dicts con keys 'name' y 'comment' para cada esquema.
    """
    if schema_filter:
        # Modo específico: solo el esquema indicado
        rows = spark.sql(f"""
            SELECT schema_name, comment
            FROM {catalog}.information_schema.schemata
            WHERE catalog_name = '{catalog}'
              AND schema_name = '{schema_filter}'
        """).collect()
    else:
        # Modo catálogo completo: todos los esquemas excepto los del sistema
        rows = spark.sql(f"""
            SELECT schema_name, comment
            FROM {catalog}.information_schema.schemata
            WHERE catalog_name = '{catalog}'
              AND schema_name NOT IN ('information_schema', 'default')
        """).collect()

    schemas = [{"name": r["schema_name"], "comment": r["comment"] or ""} for r in rows]
    logger.info(f"  Esquemas encontrados: {len(schemas)}")
    for s in schemas:
        logger.info(f"    - {s['name']}: {s['comment'][:60] if s['comment'] else '(sin comentario)'}")
    return schemas


def discover_tables_and_columns(catalog: str, schema: str) -> dict:
    """
    Descubre todas las tablas y columnas de un esquema.

    Consulta information_schema.columns con JOIN a information_schema.tables
    para obtener también el comentario existente de cada tabla.

    Returns:
        Dict indexado por nombre de tabla. Cada valor contiene:
        - comment: Comentario existente de la tabla (o cadena vacía)
        - columns: Lista de dicts con 'name' y 'type' por columna
    """
    columns_df = spark.sql(f"""
        SELECT
            c.table_name, c.column_name, c.data_type, c.ordinal_position,
            t.comment AS table_comment
        FROM {catalog}.information_schema.columns c
        LEFT JOIN {catalog}.information_schema.tables t
            ON  c.table_catalog = t.table_catalog
            AND c.table_schema  = t.table_schema
            AND c.table_name    = t.table_name
        WHERE c.table_catalog = '{catalog}'
          AND c.table_schema  = '{schema}'
        ORDER BY c.table_name, c.ordinal_position
    """).collect()

    tables = {}
    for row in columns_df:
        t = row["table_name"]
        if t not in tables:
            tables[t] = {"comment": row["table_comment"] or "", "columns": []}
        tables[t]["columns"].append({
            "name": row["column_name"],
            "type": row["data_type"],
        })

    total_cols = sum(len(t["columns"]) for t in tables.values())
    logger.info(f"    Tablas en {schema}: {len(tables)}, Columnas totales: {total_cols}")
    return tables


# -- Ejecutar descubrimiento --
schemas_to_process = discover_schemas(CATALOG_NAME, SCHEMA_NAME)

if not schemas_to_process:
    msg = (f"No se encontraron esquemas en {CATALOG_NAME}"
           + (f" con filtro '{SCHEMA_NAME}'" if SCHEMA_NAME else ""))
    logger.error(msg)
    raise ValueError(msg)

# COMMAND ----------

# MAGIC %md
# MAGIC ## 11. Ejecución Principal
# MAGIC
# MAGIC Orquesta el proceso completo de generación:
# MAGIC
# MAGIC ```
# MAGIC Para cada esquema:
# MAGIC   1. Construir contexto dinámico (insumos priorizados)
# MAGIC   2. Descubrir tablas y columnas
# MAGIC   3. Generar comentario del esquema
# MAGIC   4. Para cada tabla:
# MAGIC      a. Construir contexto dinámico (con mayor precisión)
# MAGIC      b. Obtener muestra de datos (si sampling activo)
# MAGIC      c. Generar comentario de la tabla
# MAGIC      d. Para cada columna:
# MAGIC         - Generar comentario de la columna
# MAGIC      e. Persistir todos los resultados
# MAGIC   5. Actualizar estado de la ejecución
# MAGIC ```
# MAGIC
# MAGIC El progreso se actualiza en la tabla `ejecuciones` en cada paso para
# MAGIC permitir monitoreo en tiempo real desde el dashboard.

# COMMAND ----------

# -- Generar ID único para esta ejecución --
exec_id = str(uuid.uuid4())

SEP = "=" * 60
logger.info(SEP)
logger.info(f"EJECUCIÓN INICIADA")
logger.info(f"  ID           : {exec_id}")
logger.info(f"  Timestamp    : {_now()} UTC")
logger.info(f"  Alcance      : {CATALOG_NAME}" + (f".{SCHEMA_NAME}" if SCHEMA_NAME else " (todos los esquemas)"))
logger.info(f"  Modelo       : {MODEL_ENDPOINT}")
logger.info(f"  Sampling     : {'Sí (' + str(SAMPLING_PCT) + '%)' if ENABLE_SAMPLING else 'No'}")
logger.info(SEP)

# -- Registrar inicio de ejecución --
insert_ejecucion(exec_id, "INICIADO")

try:
    # Contadores globales de progreso
    total_schemas_ok = 0
    total_tables_ok  = 0
    total_columns_ok = 0
    errors = []

    # ── Iterar sobre cada esquema a procesar ────────────────────────────────
    for schema_info in schemas_to_process:
        current_schema = schema_info["name"]
        current_schema_comment = schema_info["comment"]

        logger.info(f"\n{'─' * 50}")
        logger.info(f"ETAPA: Procesando esquema '{current_schema}'")
        logger.info(f"{'─' * 50}")

        # ── 1. Construir contexto dinámico para el esquema ──────────────────
        logger.info("  Construyendo contexto dinámico para el esquema...")
        schema_context = build_dynamic_context(
            INSUMOS, schema_name=current_schema,
        )

        # ── 2. Descubrir tablas del esquema ─────────────────────────────────
        logger.info("  Descubriendo tablas y columnas...")
        tables = discover_tables_and_columns(CATALOG_NAME, current_schema)

        if not tables:
            logger.warning(f"  Sin tablas en {current_schema} — saltando esquema")
            continue

        table_names = list(tables.keys())

        update_ejecucion(
            exec_id, "EN_PROCESO",
            f"Procesando esquema '{current_schema}' — {len(tables)} tabla(s)",
        )

        # ── 3. Generar comentario del ESQUEMA ───────────────────────────────
        logger.info("  Generando comentario del esquema...")
        try:
            generated_schema_comment = generate_schema_comment(
                schema_name=current_schema,
                tables_in_schema=table_names,
                context=schema_context,
            )
            insert_resultado(exec_id, current_schema, "__esquema__", "__esquema__", generated_schema_comment)
            total_schemas_ok += 1
            logger.info(f"  ✓ [ESQUEMA] {current_schema}: {generated_schema_comment[:100]}...")

            # Usar el comentario generado si el esquema no tenía uno
            if not current_schema_comment:
                current_schema_comment = generated_schema_comment
        except Exception as e:
            err = f"{current_schema} [esquema]: {str(e)[:200]}"
            errors.append(err)
            logger.error(f"  ✗ Error generando comentario del esquema: {e}")

        # ── 4. Iterar sobre cada tabla del esquema ──────────────────────────
        for table_name, table_data in tables.items():
            n_cols = len(table_data["columns"])
            logger.info(f"\n  TABLA: '{current_schema}.{table_name}' ({n_cols} columnas)")

            update_ejecucion(
                exec_id, "EN_PROCESO",
                f"Procesando '{current_schema}.{table_name}' — "
                f"{total_columns_ok} columnas completadas hasta ahora",
            )

            # ── 4a. Contexto dinámico para esta tabla ───────────────────────
            logger.info(f"    Construyendo contexto dinámico para la tabla...")
            table_context = build_dynamic_context(
                INSUMOS, schema_name=current_schema, table_name=table_name,
            )

            # ── 4b. Sampling de datos (si está habilitado) ──────────────────
            sample_data = ""
            if ENABLE_SAMPLING:
                logger.info(f"    Obteniendo muestra de datos...")
                sample_data = get_table_sample(
                    CATALOG_NAME, current_schema, table_name, SAMPLING_PCT,
                )

            # ── 4c. Generar comentario de la TABLA ──────────────────────────
            logger.info(f"    Generando comentario de tabla...")
            generated_table_comment = ""
            try:
                generated_table_comment = generate_table_comment(
                    schema_name=current_schema,
                    schema_comment=current_schema_comment,
                    table_name=table_name,
                    context=table_context,
                    sample_data=sample_data,
                )
                insert_resultado(exec_id, current_schema, table_name, "__tabla__", generated_table_comment)
                total_tables_ok += 1
                logger.info(f"    ✓ [TABLA] {table_name}: {generated_table_comment[:100]}...")
            except Exception as e:
                err = f"{table_name} [tabla]: {str(e)[:200]}"
                errors.append(err)
                logger.error(f"    ✗ Error en comentario de tabla: {e}")
                generated_table_comment = table_data["comment"]  # fallback al existente

            # El comentario generado (o existente) se usa como contexto para las columnas
            context_for_columns = generated_table_comment or table_data["comment"]

            # ── 4d. Generar comentarios de las COLUMNAS ─────────────────────
            logger.info(f"    Generando comentarios de {n_cols} columnas...")
            for col in table_data["columns"]:
                col_name = col["name"]
                col_type = col["type"]
                try:
                    comment = generate_column_comment(
                        schema_name=current_schema,
                        schema_comment=current_schema_comment,
                        table_name=table_name,
                        table_comment=context_for_columns,
                        column_name=col_name,
                        data_type=col_type,
                        context=table_context,
                        sample_data=sample_data,
                    )
                    insert_resultado(exec_id, current_schema, table_name, col_name, comment)
                    total_columns_ok += 1
                    logger.info(f"      ✓ {col_name} ({col_type}): {comment[:90]}...")
                except Exception as e:
                    err = f"{table_name}.{col_name}: {str(e)[:200]}"
                    errors.append(err)
                    logger.error(f"      ✗ Error en {col_name}: {e}")

    # ── Resultado final ─────────────────────────────────────────────────────
    logger.info(f"\n{SEP}")
    logger.info("ETAPA: Finalizando ejecución")

    if errors:
        estado_final = "COMPLETADO_CON_ERRORES"
        error_summary = "; ".join(errors[:5])
        resultado_final = (
            f"Completado con {len(errors)} error(es). "
            f"Esquemas: {total_schemas_ok}/{len(schemas_to_process)}. "
            f"Tablas: {total_tables_ok}. Columnas: {total_columns_ok}. "
            f"Errores: {error_summary}"
        )
    else:
        estado_final = "COMPLETADO"
        resultado_final = (
            f"Exitoso. {total_schemas_ok} esquema(s), {total_tables_ok} tabla(s) y "
            f"{total_columns_ok} columna(s) documentadas en {CATALOG_NAME}"
            + (f".{SCHEMA_NAME}" if SCHEMA_NAME else " (todos los esquemas)")
        )

    update_ejecucion(exec_id, estado_final, resultado_final)

    logger.info(f"  Estado   : {estado_final}")
    logger.info(f"  Resultado: {resultado_final}")
    logger.info(f"  Fin      : {_now()} UTC")
    logger.info(SEP)

except Exception as e:
    # Error no controlado: registrar en la tabla de ejecuciones y re-lanzar
    error_msg = str(e)[:500]
    update_ejecucion(exec_id, "ERROR", f"Error inesperado: {error_msg}")
    logger.error(f"✗ Error fatal en ejecución {exec_id}: {error_msg}")
    raise
